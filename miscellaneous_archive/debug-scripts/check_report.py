from docx import Document

doc = Document("FYP REPORT V1.docx")
text = "\n".join([p.text for p in doc.paragraphs])

checks = [
    ("Plotly mentioned", "plotly" in text.lower()),
    ("TradingView mentioned", "tradingview" in text.lower()),
    ("Next.js 14 claim", "next.js 14" in text.lower()),
    ("Next.js 16 claim", "next.js 16" in text.lower()),
    ("Walk-forward mentioned", "walk-forward" in text.lower() or "walk forward" in text.lower()),
    ("Ablation mentioned", "ablation" in text.lower()),
    ("23 features claim", "23 predictive" in text.lower()),
    ("60+ features claim", "60+" in text),
    ("89 features claim", "89" in text),
    ("WCAG mentioned", "wcag" in text.lower()),
    ("73.7% accuracy", "73.7" in text),
    ("72.5% accuracy", "72.5" in text),
    ("68% accuracy", "68%" in text),
    ("Export mentioned", "export" in text.lower()),
    ("ablation_study.py", "ablation_study" in text.lower()),
    ("Docker mentioned", "docker" in text.lower()),
    ("CI/CD mentioned", "ci/cd" in text.lower() or "github actions" in text.lower()),
    ("WebSocket mentioned", "websocket" in text.lower() or "socket.io" in text.lower()),
]

print("FYP Report V1 - Claim Verification:")
print("=" * 50)
for name, found in checks:
    icon = "PASS" if found else "MISS"
    print(f"  [{icon}] {name}")

# Count chapters
chapters = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading 1")]
print(f"\nChapters found: {len(chapters)}")
for ch in chapters:
    print(f"  - {ch}")
